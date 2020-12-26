from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from form.Utils import save_project_form_file


def generate_apply_budget_form(project_id: str, project_name: str, company: str, year: str,
                               budget_info: dict):
    doc = Document('./template-forms/3 协作成本预算表.docx')
    doc.styles['Normal'].font.name = u'仿宋_GB2312'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    company_paragraph = doc.paragraphs[1]
    company_paragraph.clear()
    run = company_paragraph.add_run('供应商名称：' + company)
    run.font.size = Pt(12)

    project_paragraph = doc.paragraphs[2]
    project_paragraph.clear()
    run = project_paragraph.add_run('项目号/项目名称：' + project_id + '/' + project_name)
    run.font.size = Pt(12)
    table = doc.tables[0]
    # 协作内容 列0 工作量 列1 单价 列2 总计 列3 合计列3
    budget_list = budget_info['budget_list']
    total_amount = budget_info['total_budget']
    run = table.rows[-1].cells[3].paragraphs[0].add_run(total_amount)
    run.font.size = Pt(12)

    r = 1
    for (content, amount, price, total) in budget_list:
        run = table.rows[r].cells[0].paragraphs[0].add_run(content)
        run.font.size = Pt(12)
        run = table.rows[r].cells[1].paragraphs[0].add_run(amount)
        run.font.size = Pt(12)
        run = table.rows[r].cells[2].paragraphs[0].add_run(price)
        run.font.size = Pt(12)
        run = table.rows[r].cells[3].paragraphs[0].add_run(total)
        run.font.size = Pt(12)
        r = r + 1
    save_project_form_file(company=company, year=year,
                           project_id=project_id, project_name=project_name,
                           form_name='3 协作成本预算表', doc=doc)
