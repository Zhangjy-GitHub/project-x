from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING

from form.Utils import save_company_form_file


def generate_quarter_confirm_form(service_company: str, year: str, quarter: str, confirm_info: dict):
    doc = Document('./template-forms/5 协作工作确认单.docx')
    doc.styles['Normal'].font.name = u'仿宋_GB2312'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')

    table = doc.tables[0]
    # 替换：(0, 1)协作商：  (1, 0)确认类型：□季度/□年度/+ u'\u2713' +单项目/□临时 
    # 替换：(2, 0)协作区间：     年     月-    年     月
    # 行4 开始：序号 列0，项目名称 列1 工作量 列2 项目号 列3 备注 列4
    table.rows[0].cells[1].paragraphs[0].clear()
    run = table.rows[0].cells[1].paragraphs[0].add_run('协作商：' + service_company)
    run.font.size = Pt(12)
    table.rows[1].cells[0].paragraphs[0].clear()
    run = table.rows[1].cells[0].paragraphs[0].add_run('确认类型：' + u'\u2713' + '季度/□年度/□单项目/□临时')
    run.font.size = Pt(12)
    table.rows[2].cells[0].paragraphs[0].clear()
    quarter_month = {'1': ('1', '3'), '2': ('4', '6'), '3': ('7', '9'), '4': ('10', '12')}
    service_time = year + '年' + quarter_month[quarter][0] + '月 —— ' + year + \
                   '年' + quarter_month[quarter][1] + '月'
    run = table.rows[2].cells[0].paragraphs[0].add_run('协作区间：' + service_time)
    run.font.size = Pt(12)

    r = 4
    for (p, confirm_list) in confirm_info.items():
        project_id_name = p.split('+', 1)
        project_id = project_id_name[0]
        project_name = project_id_name[1]
        for (service_amount, price, total, remark) in confirm_list:
            if r >= len(table.rows):
                table.add_row()
            run = table.rows[r].cells[0].paragraphs[0].add_run(project_name)
            run.font.size = Pt(12)
            run = table.rows[r].cells[1].paragraphs[0].add_run(service_amount)
            run.font.size = Pt(12)
            run = table.rows[r].cells[2].paragraphs[0].add_run(price)
            run.font.size = Pt(12)
            run = table.rows[r].cells[3].paragraphs[0].add_run(str(total))
            run.font.size = Pt(12)
            run = table.rows[r].cells[4].paragraphs[0].add_run(project_id)
            run.font.size = Pt(12)
            run = table.rows[r].cells[5].paragraphs[0].add_run(remark)
            run.font.size = Pt(12)
            r += 1

    table.add_row()

    table.rows[r].cells[0].merge(table.rows[r].cells[3])
    p = table.rows[r].cells[0].add_paragraph('委托方：辽宁北方实验室有限公司（盖章）')
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.bold = True
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    p = table.rows[r].cells[0].add_paragraph('负责人（签字）：')
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.bold = True
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    p = table.rows[r].cells[0].add_paragraph('')
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    p = table.rows[r].cells[0].add_paragraph('')
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    p = table.rows[r].cells[0].add_paragraph('    年   月   日')
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.bold = True
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    p = table.rows[r].cells[0].add_paragraph('')
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    table.rows[r].cells[4].merge(table.rows[r].cells[-1])
    content = '协作商：' + service_company + '（盖章）'
    p = table.rows[r].cells[4].add_paragraph(content)
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.bold = True
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    p = table.rows[r].cells[4].add_paragraph('负责人（签字）：')
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.bold = True
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    if len(content) <= 20:
        p = table.rows[r].cells[4].add_paragraph('')
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    p = table.rows[r].cells[4].add_paragraph('')
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    p = table.rows[r].cells[4].add_paragraph('    年   月   日')
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.bold = True
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    p = table.rows[r].cells[4].add_paragraph('')
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    save_company_form_file(company=service_company, year=year, form_name='5 协作工作确认单_季度_' + quarter, doc=doc)
    # 公司名称文件夹存在则保存，否则不保存
    #save_company_exists_form_file(company=service_company, year=year, form_name='5 协作工作确认单_季度_' + quarter, doc=doc)