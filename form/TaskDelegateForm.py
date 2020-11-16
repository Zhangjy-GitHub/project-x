from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from form.Utils import save_project_form_file


def generate_task_delegate_form(project_id: str, project_name: str, company: str, year: str, task_infos: dict):
    doc = Document('./template-forms/4 协作工作委派单.docx')
    doc.styles['Normal'].font.name = u'仿宋_GB2312'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    # 2 协作商：  3 协作项目:  4 项目执行时间：年 月 —— 年 月  5 项目地点： 7 委托方项目经理： 8  委托方联系人：  联系电话：
    doc.paragraphs[2].clear()
    run = doc.paragraphs[2].add_run('协作商：' + company)
    run.font.size = Pt(12)
    run.font.bold = True

    doc.paragraphs[3].clear()
    run = doc.paragraphs[3].add_run('协作项目: ' + project_id + '/' + project_name)
    run.font.size = Pt(12)
    run.font.bold = True

    doc.paragraphs[4].clear()
    run = doc.paragraphs[4].add_run('项目执行时间：' + task_infos['service_time'])
    run.font.size = Pt(12)
    run.font.bold = True

    doc.paragraphs[5].clear()
    run = doc.paragraphs[5].add_run('项目地点：' + task_infos['service_area'])
    run.font.size = Pt(12)
    run.font.bold = True

    doc.paragraphs[7].clear()
    run = doc.paragraphs[7].add_run('委托方项目经理：' + '隋大智')
    run.font.size = Pt(12)
    run.font.bold = True

    table = doc.tables[0]
    # 协作内容从1行开始，列1 协作内容 列2 工作量 列3 备注
    r = 1
    task_list = task_infos['task_list']
    for (content, amount, remarks) in task_list:
        run = table.rows[r].cells[0].paragraphs[0].add_run(str(r))
        run.font.size = Pt(12)

        run = table.rows[r].cells[1].paragraphs[0].add_run(content)
        run.font.size = Pt(12)

        run = table.rows[r].cells[2].paragraphs[0].add_run(amount)
        run.font.size = Pt(12)

        run = table.rows[r].cells[3].paragraphs[0].add_run(remarks)
        run.font.size = Pt(12)
        r = r + 1
    save_project_form_file(company, year, project_id, project_name, '4 协作工作委派单', doc)