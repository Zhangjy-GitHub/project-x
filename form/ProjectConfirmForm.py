from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING

from form.Utils import save_project_form_file


def generate_project_confirm_form(project_id: str, project_name: str,
                                  company: str, year: str, confirm_info: dict):
    doc = Document('./template-forms/5 协作工作确认单.docx')
    doc.styles['Normal'].font.name = u'仿宋_GB2312'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    table = doc.tables[0]
    # 替换：(0, 1)协作商：  (1, 0)确认类型：□季度/□年度/+ u'\u2713' +单项目/□临时
    # 替换：(2, 0)协作区间：     年     月-    年     月
    # 行4 开始：序号 列0，项目名称 列1 工作量 列2 项目号 列3 备注 列4
    table.rows[0].cells[1].paragraphs[0].clear()
    run = table.rows[0].cells[1].paragraphs[0].add_run('协作商：' + confirm_info['service_company'])
    run.font.size = Pt(12)
    table.rows[1].cells[0].paragraphs[0].clear()
    run = table.rows[1].cells[0].paragraphs[0].add_run('确认类型：□季度/□年度/' + u'\u2713' + '单项目/□临时')
    run.font.size = Pt(12)
    table.rows[2].cells[0].paragraphs[0].clear()
    run = table.rows[2].cells[0].paragraphs[0].add_run('协作区间：' + confirm_info['service_time'])
    run.font.size = Pt(12)
    confirm_info_list = confirm_info['confirm_list']
    run = table.rows[4].cells[0].paragraphs[0].add_run(project_name)
    run.font.size = Pt(12)
    run = table.rows[4].cells[1].paragraphs[0].add_run(confirm_info_list[0][0])
    run.font.size = Pt(12)
    run = table.rows[4].cells[2].paragraphs[0].add_run(confirm_info_list[0][1])
    run.font.size = Pt(12)
    run = table.rows[4].cells[3].paragraphs[0].add_run(str(confirm_info_list[0][2]))
    run.font.size = Pt(12)
    run = table.rows[4].cells[4].paragraphs[0].add_run(project_id)
    run.font.size = Pt(12)
    run = table.rows[4].cells[5].paragraphs[0].add_run(confirm_info_list[0][3])
    run.font.size = Pt(12)

    table.add_row()

    table.rows[5].cells[0].merge(table.rows[5].cells[3])
    p = table.rows[5].cells[0].add_paragraph('委托方：辽宁北方实验室有限公司')
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.bold = True
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    p = table.rows[5].cells[0].add_paragraph('负责人（签字）：')
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.bold = True
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    p = table.rows[5].cells[0].add_paragraph('（实施主管部门）')
    p.runs[0].font.size = Pt(12)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    p = table.rows[5].cells[0].add_paragraph('')
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    p = table.rows[5].cells[0].add_paragraph('    年   月   日')
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.bold = True
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    p = table.rows[5].cells[0].add_paragraph('')
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    table.rows[5].cells[4].merge(table.rows[5].cells[-1])
    content = '协作商：' + confirm_info['service_company']
    p = table.rows[5].cells[4].add_paragraph(content)
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.bold = True
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    p = table.rows[5].cells[4].add_paragraph('负责人（签字）：')
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.bold = True
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    if len(content) <= 19:
        p = table.rows[5].cells[4].add_paragraph('')
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    p = table.rows[5].cells[4].add_paragraph('')
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    p = table.rows[5].cells[4].add_paragraph('    年   月   日')
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.bold = True
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    p = table.rows[5].cells[4].add_paragraph('')
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    save_project_form_file(company=company, year=year, project_id=project_id,
                           project_name=project_name,
                           form_name='5 协作工作确认单', doc=doc)
