from docx import Document
from docx.shared import Pt
import os


def generate_contract(contract_info: dict):
    doc = Document('./template-forms/技术服务合同.docx')
    # 替换
    config_info = {24: 'company_name', 25: 'company_area', 26: 'legal_person', 29: 'comm_address',
                   33: 'contract_area', 111: 'company_name'}
    """
    6: '乙方（全称）：' + company_name, 24: '受托方 （乙方）：' + company_name,
    33: '甲方、乙方本着互惠互利，合作共赢，持续发展的原则，经友好协商，就' + contract_info['contract_area'] +
                       '的甲方信息化相关业务进行技术服务合作。'
                       '双方经过平等协商，在真实、充分地表达各自意愿的基础上，根据《中华人民共和国合同法》的规定，达成如下协议，并由双方共同恪守。',
    111: '乙方：' + company_name + '              （盖章）'
    """

    p = doc.paragraphs[6]
    run = p.add_run(contract_info['company_name'])
    run.font.size = Pt(16)
    run.font.name = u'宋体'
    run.font.bold = True

    for (position, key) in config_info.items():
        p = doc.paragraphs[position]
        for r in p.runs:
            if r.underline:
                origin_length = len(r.text)
                r.clear()
                r.text = contract_info[key]
                if len(r.text) < origin_length:
                    r.text += ' ' * (origin_length - len(r.text) * 2 - 1)
                r.font.size = Pt(12)
                r.font.name = '宋体'
                r.underline = True

    if not os.path.exists('../生成表格/'):
        os.mkdir('../生成表格/')
    if not os.path.exists('./生成表格/' + contract_info['company_name'] + '/'):
        os.mkdir('./生成表格/' + contract_info['company_name'] + '/')
    if not os.path.exists('./生成表格/' + contract_info['company_name'] + '/' + contract_info['contract_year'] + '/'):
        os.mkdir('./生成表格/' + contract_info['company_name'] + '/' + contract_info['contract_year'] + '/')
    doc.save('./生成表格/' + contract_info['company_name'] + '/' + contract_info['contract_year'] + '/' +
             '技术服务合同_' + contract_info['contract_year'] + '.docx')


if __name__ == '__main__':
    contract = {
        'company_area': '合肥市',
        'legal_person': '张克勤',
        'comm_address': '合肥市蜀山区长江西路与湖光路交口东北角上堤公寓乐客来国际商业中心公寓式酒店1812',
        'contract_area': '华东地区',
        'company_name': '安徽佰睿信息技术咨询有限公司',
        'contract_year': '2019'
    }
    generate_contract(contract_info=contract)
    pass
