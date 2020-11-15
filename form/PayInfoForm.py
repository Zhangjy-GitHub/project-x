from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from form import RMBConvert
from form.Utils import save_company_form_file


def generate_pay_info_form(company: str, year: str, pay_info: dict):
    doc = Document('C:\\Users\\zhang\\Documents\\Tools\\template-forms\\6 协作工作结算单.docx')
    doc.styles['Normal'].font.name = u'仿宋_GB2312'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    # 协作商 (0, 1) 
    # 替换 申请结算金额 (1, 1) 人民币：¥              （大写：          ）
    # 替换 结算描述 (2, 1) 本期工作量已审核。
    # 替换 本次结算金额 (3, 1) 人民币：¥              （大写：          ）
    table = doc.tables[0]
    run = table.rows[0].cells[1].paragraphs[0].add_run(company)
    run.font.size = Pt(12)

    table.rows[1].cells[1].paragraphs[0].clear()
    run = table.rows[1].cells[1].paragraphs[0].add_run('人民币：¥ ' + pay_info['apply_pay'] +
                                                       '  （大写：' + RMBConvert.convert(pay_info['apply_pay']) + '）')
    run.font.size = Pt(12)

    for i in range(1, len(table.rows[2].cells[1].paragraphs)):
        p = table.rows[2].cells[1].paragraphs[1]
        p._element.getparent().remove(p._element)
    table.rows[2].cells[1].paragraphs[0].clear()
    run = table.rows[2].cells[1].paragraphs[0].add_run('本期工作量已审核。')
    run.font.size = Pt(12)

    table.rows[3].cells[1].paragraphs[0].clear()
    run = table.rows[3].cells[1].paragraphs[0].add_run('人民币：¥ ' + pay_info['actual_pay']+
                                                       '  （大写：' + RMBConvert.convert(pay_info['actual_pay']) + '）')
    run.font.size = Pt(12)

    save_company_form_file(company, year, '6 协作工作结算单_' + str(pay_info['actual_pay']), doc)


if __name__ == '__main__':
    pay_info = {'apply_pay': '1000', 'actual_pay': '1000'}
    generate_pay_info_form('xxxx', '2018', pay_info=pay_info)
    pass
