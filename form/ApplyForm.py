from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from form.Utils import save_project_form_file


def generate_apply_form(project_id: str, project_name: str,
                        company: str, year: str, apply_info: dict):
    doc = Document('./template-forms/2 技术协作申请表.docx')
    doc.styles['Normal'].font.name = u'仿宋_GB2312'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')

    table = doc.tables[0]

    run = table.rows[0].cells[1].paragraphs[0].add_run(apply_info['apply_dep'])
    run.font.size = Pt(12)

    run = table.rows[0].cells[4].paragraphs[0].add_run(apply_info['apply_time'])
    run.font.size = Pt(12)

    # 替换  项目号/项目名称：\n
    table.rows[1].cells[1].text = ''
    run = table.rows[1].cells[1].paragraphs[0].add_run('项目号/项目名称：\n' + project_id +
                                                       '/' + project_name)
    run.font.size = Pt(12)

    run = table.rows[2].cells[1].paragraphs[0].add_run(company)
    run.font.size = Pt(12)

    # 替换 u'\u2713' + '□ 1）外场支持协助类\n□ 2）驻场技术支持类\n□ 3）设备验货支持类\n□ 4）项目文件评审类\n□ 5）专业技术支持类'
    table.rows[3].cells[1].text = ''
    run = table.rows[3].cells[1].paragraphs[0].add_run(apply_info['service_type'])
    run.font.size = Pt(12)

    # 替换 年   月 ——   年   月
    table.rows[4].cells[1].text = ''
    run = table.rows[4].cells[1].paragraphs[0].add_run(apply_info['service_time'])
    run.font.size = Pt(12)

    # 替换 万                协作成本构成详见附表
    table.rows[5].cells[1].text = ''
    run = table.rows[5].cells[1].paragraphs[0].add_run(apply_info['service_budget'] +
                                                       '万                协作成本构成详见附表')
    run.font.size = Pt(12)
    save_project_form_file(company, year, project_id, project_name, '2 技术协作申请表', doc)
