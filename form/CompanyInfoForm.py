from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from form.Utils import save_company_form_file


def generate_company_info_form(companies_info: dict):
    doc = Document('./template-forms/1 供应商情况.docx')
    doc.styles['Normal'].font.name = u'仿宋_GB2312'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋_GB2312')
    p = doc.paragraphs[2]
    p.clear()
    run = p.add_run('服务类型：□设备采购/□办公用品采购/' + u'\u2713' + '技术协作/□市场推广')
    run.font.size = Pt(12)
    # 协作商(0, 1)  法定代表人(0, 7) 注册地址(1, 1) 注册资金(2, 1) 成立时间(2, 4) 员工数(2, 8)
    # 经营范围(3, 1)  开户银行(4, 1) 银行账号(4, 7) 联系人(5, 1) 联系电话(5, 7)
    # 营业执照(8,2) 企查查(9,2) 信用中国(10, 2) 专业符合(11, 2)
    info_poisons = {'company': (0, 1), 'person': (0, 7), 'register_address': (1, 1),
                    'register_fund': (2, 1), 'founded_time': (2, 4), 'employee_total': (2, 8),
                    'business_scope': (3, 1), 'account_bank': (4, 1), 'account_num': (4, 7),
                    'contact_person': (5, 1), 'contact_phone': (5, 7), 'bank_name': (4, 1),
                    'bank_account': (4, 7)}
    check_poisons = [(8, 2), (9, 2), (10, 2), (11, 2)]
    table = doc.tables[0]
    for (k, p) in info_poisons.items():
        if k in companies_info.keys():
            run = table.rows[p[0]].cells[p[1]].paragraphs[0].add_run(companies_info[k])
            run.font.size = Pt(12)
    for p in check_poisons:
        pr = table.rows[p[0]].cells[p[1]].paragraphs[0]
        pr.clear()
        pr.add_run(u'\u2713' + '无。□有：')
        run.font.size = Pt(12)
    save_company_form_file(company=companies_info['company'], year=None,
                           form_name='1 供应商情况', doc=doc)


if __name__ == '__main__':
    company_info = {'company': '111'}
    generate_company_info_form(companies_info=company_info)
    pass
